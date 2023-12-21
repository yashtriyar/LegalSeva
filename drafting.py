import config
import openai
import docx
import datetime
import os

current_date = datetime.date.today()
formatted_date = current_date.strftime("%Y-%m-%d")
openai.api_key = config.OpenAI
def terms(text):
    
    
    max_characters = 3000
    text_to_process = text[:max_characters]

    # Define your prompt based on the extracted text
    prompt = f"Write terms and conditons of legal document on {text_to_process}"

    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=2000,
        temperature=0.7,
    )

    return store_terms_and_conditions(response.choices[0].text)
    

def addpr(text):
    
    # Limit the text to a maximum number of characters (adjust as needed)
    max_characters = 3000
    text_to_process = text[:max_characters]

    # Define your prompt based on the extracted text
    prompt = f"write this text '{text_to_process}' as one of the terms of contract. write it properly in 1 sentence with legal words and carefull clauses/."

    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=2000,
        temperature=0.7,
    )

    return response.choices[0].text





def store_terms_and_conditions(text):
    # Split the text into lines
    lines = text.split('\n')
    
    # Initialize an empty list to store the terms
    terms_and_conditions = []

    # Iterate through the lines, stripping the numbering and adding to the list
    for line in lines:
        # Remove the numbering and leading/trailing whitespace
        term = line.split('. ', 1)[-1].strip()
        
        # Add the term to the list if it's not empty
        if term:
            terms_and_conditions.append(term)

    return terms_and_conditions







def modification(terms_and_conditions):
    for i,a in enumerate(terms_and_conditions):
        print(f'{i+1} :{a}')
    action = input("Enter 'add' to add a condition, 'modify' to modify a condition, 'delete' to delete a condition, or 'done' to finish: ")
    while action.lower() != "done":
        if action.lower() == "modify":
            index_to_modify = int(input("Enter the number of the condition to modify: "))
            if 1 <= index_to_modify <= len(terms_and_conditions):
                new_condition = input(f"Enter the new condition for item {index_to_modify}: ")
                legalised=addpr(new_condition)
                s=legalised.strip()
                sr=f'{index_to_modify}. {s}'
                terms_and_conditions[index_to_modify - 1] = sr
                print(f"Condition {index_to_modify} modified.")
            else:
                print("Invalid condition number.")
        elif action.lower() == "delete":
            index_to_delete = int(input("Enter the number of the condition to delete: "))
            if 1 <= index_to_delete <= len(terms_and_conditions):
                del terms_and_conditions[index_to_delete - 1]
                print(f"Condition {index_to_delete} deleted.")
            else:
                print("Invalid condition number.")
        elif action.lower() == "add":
            pr=input("Enter the condition you want to add")
            proper=addpr(pr)
            terms_and_conditions.append(proper.strip())
        else:
            print("Invalid action. Enter 'modify', 'delete', or 'done'.")

        # Display the updated terms and conditions
        for a in terms_and_conditions:
            print(f'{i+1} :{a}')

        # Ask for the next action
        action = input("Enter 'add' to add a condition, 'modify' to modify a condition, 'delete' to delete a condition, or 'done' to finish: ")
    print("\nFinal Terms and Conditions:")
    for a in terms_and_conditions:
        print(a)
    return terms_and_conditions


def generate_legal_document(text,terms,inputs,dates):
    
    # Limit the text to a maximum number of characters (adjust as needed)
    max_characters = 3000
    text_to_process = text[:max_characters]

    # Define your prompt based on the extracted text
    prompt = f'''Write a legal document for '{text_to_process}' include these terms and conditions {terms} highlight the main topic and having the inputs{inputs}/.the format should be 
    Title : which type of document we are printing. highlight the headings

    date - {dates}
    details of party 1 probably who are providing or giving like employer , owner
    details of party 2 who are seeking
    place
    terms and conditons which heading and highlishts marking important clauses
    clauses
    place for signature'''

    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=2000,
        temperature=0.7,
    )

    return response.choices[0].text

def create_word_document(contract_content):
    # Get the current timestamp to create a unique filename
    current_time = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
    
    # Specify the folder name
    folder_name = "documents"

    
    folder_path = os.path.join(os.path.dirname(os.path.realpath(__file__)), folder_name)

    
    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    
    filename = os.path.join(folder_path, f"Generated_Contract_{current_time}.docx")

    doc = docx.Document()
    doc.add_paragraph(contract_content)

    doc.save(filename)
    print(f"Generated Contract saved to '{filename}'")

# def create_word_document(contract_content):
#     # Get the current timestamp to create a unique filename
#     current_time = datetime.datetime.now().strftime("%Y%m%d%H%M%S")
#     filename = f"Generated_Contract_{current_time}.docx"

#     doc = docx.Document()
#     doc.add_paragraph(contract_content)

#     doc.save(filename)
#     print(f"Generated Contract saved to '{filename}'")

def input_needs_generate(doctype):
    
    # Limit the text to a maximum number of characters (adjust as needed)
    max_characters = 3000
    #text_to_process = text[:max_characters]

    # Define your prompt based on the extracted text
    prompt = f"write the set of inputs required to draft legal documents {doctype}/."

    response = openai.Completion.create(
        engine="text-davinci-003",
        prompt=prompt,
        max_tokens=2000,
        temperature=0.7,
    )

    return store_terms_and_conditions(response.choices[0].text)


def get_user_input(questions):
    user_responses = []  # Create an empty list to store user responses
    
    for i, question in enumerate(questions, start=1):
        user_response = input(f"Q{i}: {question} - ")
        user_responses.append(f"Q{i}: {user_response}")
    
    return '\n'.join(user_responses)


def main():
    doctyp=input("Type of Legal Document(eg. Rental Agreement)")
    inputs=input_needs_generate(doctyp)
    ques=get_user_input(inputs)
    term=terms(doctyp)
    modified_terms=modification(term)
    current_date = datetime.date.today()
    formatted_date = current_date.strftime("%Y-%m-%d")
    #terms_arrays=store_terms_and_conditions(modified_terms)
    doc=generate_legal_document(doctyp,modified_terms,ques,formatted_date)
    create_word_document(doc)
    print(doc)
main()
    

#first input will be given and stored
#terms and conditions will be generated and given
#it will be modified and added or edited
#the code to generate the legal document based on terms and conditions and other parameter



































